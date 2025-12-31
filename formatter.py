# formatter.py
import os
from dataclasses import dataclass, asdict
from typing import Optional, Dict, Any, List, Tuple
import json
import io

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ALIGN_MAP = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
}

PAGE_FMT_MAP = {
    "DECIMAL": "decimal",
    "ROMAN_UPPER": "upperRoman",
    "ROMAN_LOWER": "lowerRoman",
    "LETTER_UPPER": "upperLetter",
    "LETTER_LOWER": "lowerLetter",
}

PAGE_NUMBER_POSITION = [
    "FOOTER_LEFT", "FOOTER_CENTER", "FOOTER_RIGHT",
    "HEADER_LEFT", "HEADER_CENTER", "HEADER_RIGHT",
]

PAPER_PRESET = {
    "A4_PORTRAIT": {"width_cm": 21.0, "height_cm": 29.7},
    "A4_LANDSCAPE": {"width_cm": 29.7, "height_cm": 21.0},
}

@dataclass
class StyleConfig:
    font_name: str = "Times New Roman"
    font_size_pt: float = 13.0
    bold: bool = False
    italic: bool = False
    color_hex: Optional[str] = None
    line_spacing: float = 1.5
    space_before_pt: float = 0.0
    space_after_pt: float = 6.0
    first_line_indent_cm: float = 1.0
    alignment: str = "JUSTIFY"
    keep_with_next: bool = False
    page_break_before: bool = False

@dataclass
class PageSetupConfig:
    paper: str = "A4_PORTRAIT"
    margin_left_cm: float = 3.5
    margin_right_cm: float = 2.0
    margin_top_cm: float = 2.0
    margin_bottom_cm: float = 2.0
    header_distance_cm: float = 1.25
    footer_distance_cm: float = 1.25
    different_first_page: bool = False

@dataclass
class PageNumberConfig:
    enabled: bool = True
    position: str = "FOOTER_CENTER"
    template: str = "Trang {PAGE}/{NUMPAGES}"
    start_at: int = 1
    restart_each_section: bool = False
    number_format: str = "DECIMAL"
    font_name: str = "Times New Roman"
    font_size_pt: float = 11.0

@dataclass
class TocConfig:
    insert_toc: bool = False
    heading_levels: str = "1-3"
    title: str = "MỤC LỤC"
    title_bold: bool = True
    title_font_size_pt: float = 14.0
    title_alignment: str = "CENTER"

@dataclass
class ProcessingConfig:
    force_run_font_everywhere: bool = True
    force_paragraph_format_everywhere: bool = True
    include_tables: bool = True

@dataclass
class ReportConfig:
    normal: StyleConfig = StyleConfig()
    title: StyleConfig = StyleConfig(font_size_pt=16.0, bold=True, alignment="CENTER", space_after_pt=12.0, first_line_indent_cm=0.0, line_spacing=1.2)
    heading1: StyleConfig = StyleConfig(font_size_pt=14.0, bold=True, alignment="LEFT", space_before_pt=12.0, space_after_pt=6.0, first_line_indent_cm=0.0, line_spacing=1.2, keep_with_next=True)
    heading2: StyleConfig = StyleConfig(font_size_pt=13.0, bold=True, alignment="LEFT", space_before_pt=10.0, space_after_pt=4.0, first_line_indent_cm=0.0, line_spacing=1.2, keep_with_next=True)
    heading3: StyleConfig = StyleConfig(font_size_pt=13.0, bold=True, italic=True, alignment="LEFT", space_before_pt=8.0, space_after_pt=4.0, first_line_indent_cm=0.0, line_spacing=1.2, keep_with_next=True)
    caption: StyleConfig = StyleConfig(font_size_pt=11.0, italic=True, alignment="CENTER", space_before_pt=6.0, space_after_pt=6.0, first_line_indent_cm=0.0, line_spacing=1.0)
    pagesetup: PageSetupConfig = PageSetupConfig()
    pagenumber: PageNumberConfig = PageNumberConfig()
    toc: TocConfig = TocConfig()
    processing: ProcessingConfig = ProcessingConfig()

def cfg_to_dict(cfg: ReportConfig) -> Dict[str, Any]:
    return asdict(cfg)

def cfg_from_dict(d: Dict[str, Any]) -> ReportConfig:
    default = cfg_to_dict(ReportConfig())
    merged = deep_merge(default, d)

    def build_style(x): return StyleConfig(**x)
    def build_pagesetup(x): return PageSetupConfig(**x)
    def build_pagenumber(x): return PageNumberConfig(**x)
    def build_toc(x): return TocConfig(**x)
    def build_processing(x): return ProcessingConfig(**x)

    return ReportConfig(
        normal=build_style(merged["normal"]),
        title=build_style(merged["title"]),
        heading1=build_style(merged["heading1"]),
        heading2=build_style(merged["heading2"]),
        heading3=build_style(merged["heading3"]),
        caption=build_style(merged["caption"]),
        pagesetup=build_pagesetup(merged["pagesetup"]),
        pagenumber=build_pagenumber(merged["pagenumber"]),
        toc=build_toc(merged["toc"]),
        processing=build_processing(merged["processing"]),
    )

def deep_merge(a: Dict[str, Any], b: Dict[str, Any]) -> Dict[str, Any]:
    out = dict(a)
    for k, v in (b or {}).items():
        if isinstance(v, dict) and isinstance(out.get(k), dict):
            out[k] = deep_merge(out[k], v)
        else:
            out[k] = v
    return out

def load_config_json_bytes(data: bytes) -> ReportConfig:
    d = json.loads(data.decode("utf-8"))
    return cfg_from_dict(d)

def save_config_json_bytes(cfg: ReportConfig) -> bytes:
    return json.dumps(cfg_to_dict(cfg), ensure_ascii=False, indent=2).encode("utf-8")

def _set_font_all_scripts(font_element, font_name: str):
    rFonts = font_element.rFonts
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)

def _apply_style_to_docx_style(docx_style, cfg: StyleConfig):
    docx_style.font.name = cfg.font_name
    _set_font_all_scripts(docx_style.font, cfg.font_name)
    docx_style.font.size = Pt(cfg.font_size_pt)
    docx_style.font.bold = cfg.bold
    docx_style.font.italic = cfg.italic
    if cfg.color_hex:
        try:
            docx_style.font.color.rgb = bytes.fromhex(cfg.color_hex)
        except Exception:
            pass

    pf = docx_style.paragraph_format
    pf.line_spacing = cfg.line_spacing
    pf.space_before = Pt(cfg.space_before_pt)
    pf.space_after = Pt(cfg.space_after_pt)
    pf.first_line_indent = Cm(cfg.first_line_indent_cm) if cfg.first_line_indent_cm else None
    pf.alignment = ALIGN_MAP.get(cfg.alignment.upper(), WD_ALIGN_PARAGRAPH.JUSTIFY)
    pf.keep_with_next = cfg.keep_with_next
    pf.page_break_before = cfg.page_break_before

def _iter_all_paragraphs(doc: Document, include_tables: bool = True):
    for p in doc.paragraphs:
        yield p
    if include_tables:
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

def _add_simple_field(paragraph, instr: str):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)

def _clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)

def _set_section_page_numbering(section, start_at: int, fmt: str, restart: bool):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)

    fmt_val = PAGE_FMT_MAP.get(fmt.upper(), "decimal")
    pgNumType.set(qn("w:fmt"), fmt_val)
    if restart:
        pgNumType.set(qn("w:start"), str(int(start_at)))
    else:
        if pgNumType.get(qn("w:start")) is not None:
            pgNumType.attrib.pop(qn("w:start"), None)

class DocxReportFormatter:
    def __init__(self, config: ReportConfig):
        self.cfg = config

    def format_docx_bytes(self, input_docx_bytes: bytes) -> bytes:
        doc = Document(io.BytesIO(input_docx_bytes))

        self._apply_page_setup(doc)
        self._apply_styles(doc)
        self._insert_toc_if_needed(doc)
        self._apply_page_numbers(doc)
        self._force_paragraph_format(doc)
        self._force_run_font(doc)

        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()

    def _apply_page_setup(self, doc: Document):
        ps = self.cfg.pagesetup
        paper = PAPER_PRESET.get(ps.paper, PAPER_PRESET["A4_PORTRAIT"])
        for section in doc.sections:
            section.page_width = Cm(paper["width_cm"])
            section.page_height = Cm(paper["height_cm"])
            section.left_margin = Cm(ps.margin_left_cm)
            section.right_margin = Cm(ps.margin_right_cm)
            section.top_margin = Cm(ps.margin_top_cm)
            section.bottom_margin = Cm(ps.margin_bottom_cm)
            section.header_distance = Cm(ps.header_distance_cm)
            section.footer_distance = Cm(ps.footer_distance_cm)
            section.different_first_page_header_footer = bool(ps.different_first_page)

    def _apply_styles(self, doc: Document):
        styles = doc.styles
        if "Normal" in styles:
            _apply_style_to_docx_style(styles["Normal"], self.cfg.normal)
        if "Title" in styles:
            _apply_style_to_docx_style(styles["Title"], self.cfg.title)
        for name, scfg in [("Heading 1", self.cfg.heading1), ("Heading 2", self.cfg.heading2), ("Heading 3", self.cfg.heading3)]:
            if name in styles:
                _apply_style_to_docx_style(styles[name], scfg)
        if "Caption" in styles:
            _apply_style_to_docx_style(styles["Caption"], self.cfg.caption)

    def _force_paragraph_format(self, doc: Document):
        pcfg = self.cfg.processing
        if not pcfg.force_paragraph_format_everywhere:
            return

        for p in _iter_all_paragraphs(doc, include_tables=pcfg.include_tables):
            style_name = (p.style.name if p.style is not None else "Normal").lower()

            if "heading 1" in style_name:
                scfg = self.cfg.heading1
            elif "heading 2" in style_name:
                scfg = self.cfg.heading2
            elif "heading 3" in style_name:
                scfg = self.cfg.heading3
            elif style_name == "title":
                scfg = self.cfg.title
            elif "caption" in style_name:
                scfg = self.cfg.caption
            else:
                scfg = self.cfg.normal

            pf = p.paragraph_format
            pf.line_spacing = scfg.line_spacing
            pf.space_before = Pt(scfg.space_before_pt)
            pf.space_after = Pt(scfg.space_after_pt)
            pf.first_line_indent = Cm(scfg.first_line_indent_cm) if scfg.first_line_indent_cm else None
            pf.alignment = ALIGN_MAP.get(scfg.alignment.upper(), WD_ALIGN_PARAGRAPH.JUSTIFY)
            pf.keep_with_next = scfg.keep_with_next
            pf.page_break_before = scfg.page_break_before

    def _force_run_font(self, doc: Document):
        pcfg = self.cfg.processing
        if not pcfg.force_run_font_everywhere:
            return

        for p in _iter_all_paragraphs(doc, include_tables=pcfg.include_tables):
            style_name = (p.style.name if p.style is not None else "Normal").lower()

            if "heading 1" in style_name:
                scfg = self.cfg.heading1
            elif "heading 2" in style_name:
                scfg = self.cfg.heading2
            elif "heading 3" in style_name:
                scfg = self.cfg.heading3
            elif style_name == "title":
                scfg = self.cfg.title
            elif "caption" in style_name:
                scfg = self.cfg.caption
            else:
                scfg = self.cfg.normal

            for r in p.runs:
                r.font.name = scfg.font_name
                try:
                    rFonts = r._r.get_or_add_rPr().get_or_add_rFonts()
                    rFonts.set(qn("w:ascii"), scfg.font_name)
                    rFonts.set(qn("w:hAnsi"), scfg.font_name)
                    rFonts.set(qn("w:eastAsia"), scfg.font_name)
                    rFonts.set(qn("w:cs"), scfg.font_name)
                except Exception:
                    pass
                r.font.size = Pt(scfg.font_size_pt)

    def _insert_toc_if_needed(self, doc: Document):
        tc = self.cfg.toc
        if not tc.insert_toc:
            return

        p_title = doc.paragraphs[0].insert_paragraph_before(tc.title) if doc.paragraphs else doc.add_paragraph(tc.title)
        p_title.alignment = ALIGN_MAP.get(tc.title_alignment.upper(), WD_ALIGN_PARAGRAPH.CENTER)
        for run in p_title.runs:
            run.bold = tc.title_bold
            run.font.size = Pt(tc.title_font_size_pt)
            run.font.name = self.cfg.normal.font_name

        p_toc = p_title.insert_paragraph_after("")
        instr = f'TOC \\o "{tc.heading_levels}" \\h \\z \\u'
        _add_simple_field(p_toc, instr)

        p_break = p_toc.insert_paragraph_after("")
        p_break.runs[0].add_break()

    def _apply_page_numbers(self, doc: Document):
        pn = self.cfg.pagenumber
        if not pn.enabled:
            return

        for idx, section in enumerate(doc.sections):
            restart = bool(pn.restart_each_section)
            if idx == 0:
                _set_section_page_numbering(section, pn.start_at, pn.number_format, restart=True)
            else:
                _set_section_page_numbering(section, pn.start_at, pn.number_format, restart=restart)

            target_is_footer = pn.position.startswith("FOOTER")
            align_key = pn.position.split("_", 1)[1]
            alignment = {"LEFT": WD_ALIGN_PARAGRAPH.LEFT, "CENTER": WD_ALIGN_PARAGRAPH.CENTER, "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT}.get(align_key, WD_ALIGN_PARAGRAPH.CENTER)

            hf = section.footer if target_is_footer else section.header
            para = hf.paragraphs[0] if hf.paragraphs else hf.add_paragraph()

            _clear_paragraph(para)
            para.alignment = alignment

            for part in self._split_template(pn.template):
                if part == "{PAGE}":
                    _add_simple_field(para, "PAGE")
                elif part == "{NUMPAGES}":
                    _add_simple_field(para, "NUMPAGES")
                else:
                    run = para.add_run(part)
                    run.font.name = pn.font_name
                    run.font.size = Pt(pn.font_size_pt)
                    try:
                        rFonts = run._r.get_or_add_rPr().get_or_add_rFonts()
                        rFonts.set(qn("w:ascii"), pn.font_name)
                        rFonts.set(qn("w:hAnsi"), pn.font_name)
                        rFonts.set(qn("w:eastAsia"), pn.font_name)
                        rFonts.set(qn("w:cs"), pn.font_name)
                    except Exception:
                        pass

    @staticmethod
    def _split_template(template: str) -> List[str]:
        tokens = []
        i = 0
        while i < len(template):
            if template.startswith("{PAGE}", i):
                tokens.append("{PAGE}")
                i += len("{PAGE}")
            elif template.startswith("{NUMPAGES}", i):
                tokens.append("{NUMPAGES}")
                i += len("{NUMPAGES}")
            else:
                j = i
                while j < len(template) and not template.startswith("{PAGE}", j) and not template.startswith("{NUMPAGES}", j):
                    j += 1
                tokens.append(template[i:j])
                i = j
        return tokens
